#!/usr/bin/env python3
"""Write annotations into a .docx file as comments or append a notes section."""

import argparse
import json
import sys
from pathlib import Path


def write_notes_section(docx_path: str, notes: str, output_path: str | None = None) -> str:
    """Append an 'Annotations' section to the end of a .docx file."""
    try:
        from docx import Document
    except ImportError:
        sys.stderr.write(
            "python-docx is required. Install with: pip install python-docx\n"
        )
        raise SystemExit(2)

    doc = Document(docx_path)
    doc.add_page_break()
    doc.add_heading("Annotations / 阅读笔记", level=1)

    for paragraph in notes.split("\n"):
        stripped = paragraph.strip()
        if stripped:
            doc.add_paragraph(stripped)

    out = output_path or docx_path
    doc.save(out)
    return out


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(
        description="Add annotation notes to a .docx file."
    )
    parser.add_argument("--input", required=True, type=Path, help="Input .docx file")
    parser.add_argument("--output", type=Path, help="Output .docx (default: overwrite input)")
    parser.add_argument("--notes", required=True, type=Path, help="JSON or text file with notes")
    args = parser.parse_args(argv)

    if not args.input.exists():
        sys.stderr.write(f"Input not found: {args.input}\n")
        return 1

    notes_text = args.notes.read_text(encoding="utf-8")
    if args.notes.suffix == ".json":
        data = json.loads(notes_text)
        notes_text = data.get("notes", data.get("content", notes_text))

    out = write_notes_section(str(args.input), notes_text, str(args.output) if args.output else None)
    print(f"saved: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
